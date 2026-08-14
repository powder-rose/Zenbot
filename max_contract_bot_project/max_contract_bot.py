"""MAX-бот для автоматического формирования проектов договоров.

Что умеет бот:
1. Ограничивает доступ списком сотрудников.
2. По шагам спрашивает фамилию сотрудника, стоимость услуги и число комплектов.
3. Получает реквизиты текстом, изображением, DOCX или PDF.
4. Извлекает текст локально, а сканы распознаёт через Yandex Vision OCR.
5. Приводит реквизиты к единой структуре через YandexGPT.
6. Заполняет шаблон Word и отправляет готовый договор пользователю в MAX.

Для первого запуска заполните файл .env и положите рядом template2.docx.
"""

from __future__ import annotations

import base64
import json
import logging
import mimetypes
import os
import re
import threading
import time
import zipfile
from dataclasses import asdict, dataclass, field
from datetime import date, datetime
from io import BytesIO
from logging.handlers import RotatingFileHandler
from pathlib import Path
from typing import Any, Iterable
from urllib.parse import parse_qs, unquote, urlparse
from uuid import uuid4

import fitz
import requests
import truststore
from dateutil.relativedelta import relativedelta
from dotenv import load_dotenv
from docx import Document
from docxtpl import DocxTemplate
from num2words import num2words
from PIL import Image, ImageOps, UnidentifiedImageError

# На Windows это помогает Python использовать системное хранилище сертификатов.
truststore.inject_into_ssl()

# ---------------------------------------------------------------------------
# Пути и загрузка настроек
# ---------------------------------------------------------------------------

BASE_DIR = Path(__file__).resolve().parent
load_dotenv(BASE_DIR / ".env")


class ConfigurationError(RuntimeError):
    """Не хватает обязательной настройки или файла проекта."""


class MaxAPIError(RuntimeError):
    """MAX API вернул ошибку."""


class MaxAuthenticationError(MaxAPIError):
    """MAX отклонил токен бота."""


class DocumentExtractionError(RuntimeError):
    """Не удалось получить пригодный текст из файла или изображения."""


class AIServiceError(RuntimeError):
    """Не удалось получить корректный ответ от сервиса Yandex AI."""


class MissingCompanyDetailsError(ValueError):
    """В реквизитах отсутствуют обязательные поля."""

    def __init__(self, fields: Iterable[str]) -> None:
        self.fields = tuple(fields)
        super().__init__("Не хватает реквизитов: " + ", ".join(self.fields))


def require_env(name: str) -> str:
    """Возвращает обязательную переменную окружения или завершает запуск."""
    value = os.getenv(name, "").strip()
    if not value:
        raise ConfigurationError(f"Не задана переменная окружения {name}")
    return value


def parse_id_set(raw: str, variable_name: str) -> frozenset[int]:
    """Преобразует строку `123,456` в неизменяемое множество ID."""
    try:
        values = frozenset(int(part.strip()) for part in raw.split(",") if part.strip())
    except ValueError as exc:
        raise ConfigurationError(
            f"{variable_name} должен содержать числовые ID через запятую"
        ) from exc
    if not values:
        raise ConfigurationError(f"Список {variable_name} пуст")
    return values


def resolve_project_path(raw: str, default: str) -> Path:
    """Разрешает относительный путь относительно папки проекта."""
    value = Path(raw.strip() or default)
    return value if value.is_absolute() else BASE_DIR / value


@dataclass(frozen=True, slots=True)
class Config:
    """Все настройки приложения в одном объекте."""

    max_token: str
    staff_ids: frozenset[int]
    yandex_folder_id: str
    yandex_api_key: str
    yandex_model_uri: str
    max_api_url: str
    template_path: Path
    work_dir: Path
    log_dir: Path
    hello_file: Path
    log_level: str
    print_price_per_set: int
    delivery_price: int

    @classmethod
    def load(cls) -> "Config":
        token = require_env("MAX_BOT_TOKEN")
        if token.casefold().startswith("bearer "):
            raise ConfigurationError(
                "MAX_BOT_TOKEN должен содержать только токен, без слова Bearer"
            )

        folder_id = require_env("YANDEX_FOLDER_ID")
        template_path = resolve_project_path(
            os.getenv("CONTRACT_TEMPLATE", ""), "template2.docx"
        )
        if not template_path.is_file():
            raise ConfigurationError(f"Не найден шаблон договора: {template_path}")

        work_dir = resolve_project_path(os.getenv("WORK_DIR", ""), "work")
        log_dir = resolve_project_path(os.getenv("LOG_DIR", ""), "logs")
        hello_file = resolve_project_path(
            os.getenv("HELLO_FILE", ""), "messages/hello.txt"
        )
        work_dir.mkdir(parents=True, exist_ok=True)
        log_dir.mkdir(parents=True, exist_ok=True)

        def positive_int(name: str, default: int) -> int:
            raw = os.getenv(name, str(default)).strip()
            try:
                value = int(raw)
            except ValueError as exc:
                raise ConfigurationError(f"{name} должен быть целым числом") from exc
            if value < 0:
                raise ConfigurationError(f"{name} не может быть отрицательным")
            return value

        return cls(
            max_token=token,
            staff_ids=parse_id_set(require_env("MAX_STAFF_IDS"), "MAX_STAFF_IDS"),
            yandex_folder_id=folder_id,
            yandex_api_key=require_env("YANDEX_API_KEY"),
            yandex_model_uri=os.getenv(
                "YANDEX_MODEL_URI",
                f"gpt://{folder_id}/yandexgpt/latest",
            ).strip(),
            max_api_url=os.getenv(
                "MAX_API_URL", "https://platform-api2.max.ru"
            ).rstrip("/"),
            template_path=template_path,
            work_dir=work_dir,
            log_dir=log_dir,
            hello_file=hello_file,
            log_level=os.getenv("LOG_LEVEL", "INFO").upper(),
            print_price_per_set=positive_int("PRINT_PRICE_PER_SET", 3900),
            delivery_price=positive_int("DELIVERY_PRICE", 1000),
        )


def configure_logging(config: Config) -> None:
    """Пишет логи одновременно в терминал и вращаемый файл."""
    level = getattr(logging, config.log_level, logging.INFO)
    formatter = logging.Formatter(
        "%(asctime)s | %(levelname)s | %(name)s | %(message)s",
        datefmt="%Y-%m-%d %H:%M:%S",
    )

    root = logging.getLogger()
    root.setLevel(level)
    root.handlers.clear()

    console = logging.StreamHandler()
    console.setFormatter(formatter)
    root.addHandler(console)

    file_handler = RotatingFileHandler(
        config.log_dir / "max_contract_bot.log",
        maxBytes=5 * 1024 * 1024,
        backupCount=5,
        encoding="utf-8",
    )
    file_handler.setFormatter(formatter)
    root.addHandler(file_handler)


# ---------------------------------------------------------------------------
# MAX API
# ---------------------------------------------------------------------------


class MaxClient:
    """Минимальный клиент MAX API без сторонней библиотеки."""

    def __init__(self, config: Config) -> None:
        self.base_url = config.max_api_url
        self.session = requests.Session()
        self.session.trust_env = False
        self.session.headers.update(
            {
                "Authorization": config.max_token,
                "Accept": "application/json",
            }
        )
        self.upload_session = requests.Session()
        self.upload_session.trust_env = False

        # MAX ограничивает отправку двумя сообщениями в секунду в один диалог.
        self._send_locks: dict[int, threading.Lock] = {}
        self._send_locks_guard = threading.Lock()
        self._last_send_at: dict[int, float] = {}

    def _lock_for_user(self, user_id: int) -> threading.Lock:
        with self._send_locks_guard:
            return self._send_locks.setdefault(user_id, threading.Lock())

    def request(
        self,
        method: str,
        path: str,
        *,
        timeout: int | tuple[int, int] = 40,
        attempts: int = 3,
        **kwargs: Any,
    ) -> dict[str, Any]:
        """Выполняет запрос с повторами для 429 и временных ошибок сервера."""
        last_error: Exception | None = None

        for attempt in range(1, attempts + 1):
            try:
                response = self.session.request(
                    method,
                    f"{self.base_url}{path}",
                    timeout=timeout,
                    **kwargs,
                )

                if response.status_code == 401:
                    raise MaxAuthenticationError(
                        "MAX отклонил токен. Получите новый MAX_BOT_TOKEN."
                    )

                if response.ok:
                    return response.json() if response.content else {}

                detail = response.text[:800].replace("\r", " ").replace("\n", " ")
                logging.getLogger("max-api").warning(
                    "MAX API error | method=%s | path=%s | status=%s | body=%s",
                    method,
                    path,
                    response.status_code,
                    detail,
                )

                if response.status_code == 429 or response.status_code >= 500:
                    response.raise_for_status()

                raise MaxAPIError(
                    f"MAX API вернул HTTP {response.status_code}: {detail}"
                )

            except MaxAuthenticationError:
                raise
            except (requests.Timeout, requests.ConnectionError, requests.HTTPError) as exc:
                last_error = exc
                if attempt >= attempts:
                    break
                time.sleep(min(2 ** (attempt - 1), 8))
            except requests.RequestException as exc:
                raise MaxAPIError(f"Ошибка соединения с MAX: {exc}") from exc

        raise MaxAPIError("MAX API временно недоступен") from last_error

    def me(self) -> dict[str, Any]:
        return self.request("GET", "/me")

    def updates(self, marker: int | None) -> dict[str, Any]:
        params: dict[str, Any] = {
            "limit": 100,
            "timeout": 30,
            "types": "bot_started,message_created,message_callback",
        }
        if marker is not None:
            params["marker"] = marker
        return self.request(
            "GET",
            "/updates",
            params=params,
            timeout=40,
            attempts=1,
        )

    def send(
        self,
        user_id: int,
        text: str,
        attachments: list[dict[str, Any]] | None = None,
    ) -> dict[str, Any]:
        """Отправляет сообщение, соблюдая лимит на один диалог."""
        with self._lock_for_user(user_id):
            elapsed = time.monotonic() - self._last_send_at.get(user_id, 0.0)
            if elapsed < 0.55:
                time.sleep(0.55 - elapsed)

            body: dict[str, Any] = {"text": text, "format": "html"}
            if attachments:
                body["attachments"] = attachments

            result = self.request(
                "POST",
                "/messages",
                params={"user_id": user_id},
                json=body,
            )
            self._last_send_at[user_id] = time.monotonic()
            return result

    @staticmethod
    def _find_token(value: Any) -> str | None:
        """Ищет поле token в ответах разных версий загрузочного API."""
        if isinstance(value, dict):
            token = value.get("token")
            if isinstance(token, str) and token:
                return token
            for nested in value.values():
                found = MaxClient._find_token(nested)
                if found:
                    return found
        elif isinstance(value, list):
            for nested in value:
                found = MaxClient._find_token(nested)
                if found:
                    return found
        return None

    def upload_file(self, path: Path) -> dict[str, Any]:
        """Загружает файл и возвращает готовое вложение для POST /messages."""
        ticket = self.request("POST", "/uploads", params={"type": "file"})
        upload_url = str(ticket.get("url") or "")
        if not upload_url:
            raise MaxAPIError("MAX не вернул URL для загрузки файла")

        mime_type = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        with path.open("rb") as stream:
            response = self.upload_session.post(
                upload_url,
                files={"data": (path.name, stream, mime_type)},
                timeout=180,
            )
        response.raise_for_status()

        try:
            uploaded = response.json() if response.content else {}
        except ValueError as exc:
            raise MaxAPIError("Сервер загрузки MAX вернул не JSON") from exc

        token = self._find_token(uploaded) or self._find_token(ticket)
        if not token:
            # В некоторых ответах токен находится в query-параметре URL.
            query = parse_qs(urlparse(upload_url).query)
            token = next(iter(query.get("token", [])), None)
        if not token:
            raise MaxAPIError("MAX не вернул token загруженного файла")

        return {"type": "file", "payload": {"token": token}}


# ---------------------------------------------------------------------------
# YandexGPT и Yandex Vision OCR
# ---------------------------------------------------------------------------


@dataclass(slots=True)
class CompanyDetails:
    """Единая структура реквизитов, которую заполняет YandexGPT."""

    organization_type: str = ""
    organization_name: str = ""
    director_title_genitive: str = ""
    director_title_nominative: str = ""
    director_name_genitive: str = ""
    director_name_nominative: str = ""
    inn: str = ""
    legal_address: str = ""
    bank_name: str = ""
    checking_account: str = ""
    correspondent_account: str = ""
    bik: str = ""
    kpp: str = ""
    ogrn: str = ""

    LABELS = {
        "organization_type": "организационно-правовая форма",
        "organization_name": "наименование организации",
        "director_title_genitive": "должность руководителя в родительном падеже",
        "director_title_nominative": "должность руководителя",
        "director_name_genitive": "ФИО руководителя в родительном падеже",
        "director_name_nominative": "ФИО руководителя",
        "inn": "ИНН",
        "legal_address": "юридический адрес",
        "bank_name": "наименование банка",
        "checking_account": "расчётный счёт",
        "correspondent_account": "корреспондентский счёт",
        "bik": "БИК",
        "kpp": "КПП",
        "ogrn": "ОГРН или ОГРНИП",
    }
    EMPTY_VALUES = {
        "",
        "-",
        "—",
        "нет",
        "нет данных",
        "не найдено",
        "неизвестно",
        "не указано",
        "отсутствует",
        "none",
        "null",
        "n/a",
    }

    @classmethod
    def from_dict(cls, data: dict[str, Any]) -> "CompanyDetails":
        valid = {key: str(data.get(key) or "").strip() for key in cls.LABELS}
        return cls(**valid)

    @staticmethod
    def _clean(value: str) -> str:
        return " ".join(str(value or "").split())

    @property
    def is_individual_entrepreneur(self) -> bool:
        normalized = self._clean(self.organization_type).casefold()
        full_name = self._clean(self.organization_name).casefold()
        return (
            normalized == "ип"
            or "индивидуальный предприниматель" in normalized
            or "индивидуальный предприниматель" in full_name
        )

    @property
    def display_name(self) -> str:
        org_type = self._clean(self.organization_type)
        name = self._clean(self.organization_name)
        if self.is_individual_entrepreneur:
            if "индивидуальный предприниматель" in name.casefold():
                return name
            return f"Индивидуальный предприниматель {name}".strip()
        if name.casefold().startswith(org_type.casefold()):
            return name
        return f"{org_type} «{name}»".strip()

    def validate(self) -> None:
        required = set(self.LABELS)
        if self.is_individual_entrepreneur:
            required.discard("kpp")

        missing = [
            self.LABELS[key]
            for key in self.LABELS
            if key in required
            and self._clean(getattr(self, key)).casefold() in self.EMPTY_VALUES
        ]
        if missing:
            raise MissingCompanyDetailsError(missing)


class YandexAI:
    """Работа с YandexGPT и Vision OCR через REST API."""

    GPT_URL = "https://ai.api.cloud.yandex.net/foundationModels/v1/completion"
    OCR_URL = "https://ocr.api.cloud.yandex.net/ocr/v1/recognizeText"

    EXTRACTION_PROMPT = """
Ты извлекаешь реквизиты российского юридического лица или индивидуального
предпринимателя из исходного текста. Верни ТОЛЬКО корректный JSON без Markdown.
Не придумывай данные. Если значения нет, верни пустую строку.

Строго используй ключи:
{
  "organization_type": "ООО, АО, ПАО, ИП и т.п.",
  "organization_name": "наименование без кавычек и без формы, для ИП — ФИО",
  "director_title_genitive": "например: генерального директора",
  "director_title_nominative": "например: генеральный директор",
  "director_name_genitive": "ФИО в родительном падеже",
  "director_name_nominative": "ФИО в именительном падеже",
  "inn": "только цифры",
  "legal_address": "полный юридический адрес",
  "bank_name": "полное наименование банка",
  "checking_account": "20 цифр",
  "correspondent_account": "20 цифр",
  "bik": "9 цифр",
  "kpp": "9 цифр или пусто для ИП",
  "ogrn": "13 цифр или ОГРНИП 15 цифр"
}

Для ИП в полях должности укажи соответственно:
"индивидуального предпринимателя" и "индивидуальный предприниматель".
""".strip()

    def __init__(self, config: Config) -> None:
        self.folder_id = config.yandex_folder_id
        self.api_key = config.yandex_api_key
        self.model_uri = config.yandex_model_uri
        self.http = requests.Session()
        self.http.trust_env = False

    def _post_json(
        self,
        url: str,
        payload: dict[str, Any],
        *,
        operation: str,
        timeout: tuple[int, int],
        attempts: int = 3,
    ) -> dict[str, Any]:
        headers = {
            "Authorization": f"Api-Key {self.api_key}",
            "Content-Type": "application/json",
            "x-folder-id": self.folder_id,
        }
        if operation == "ocr":
            headers["x-data-logging-enabled"] = "false"

        last_error: Exception | None = None
        for attempt in range(1, attempts + 1):
            try:
                response = self.http.post(
                    url,
                    headers=headers,
                    json=payload,
                    timeout=timeout,
                )
                if response.ok:
                    return response.json()

                detail = response.text[:1000]
                if response.status_code in {429, 500, 502, 503, 504}:
                    response.raise_for_status()
                raise AIServiceError(
                    f"Yandex {operation} вернул HTTP {response.status_code}: {detail}"
                )
            except AIServiceError:
                raise
            except (requests.Timeout, requests.ConnectionError, requests.HTTPError) as exc:
                last_error = exc
                if attempt >= attempts:
                    break
                time.sleep(min(2 ** (attempt - 1), 8))
            except requests.RequestException as exc:
                raise AIServiceError(f"Ошибка соединения с Yandex {operation}") from exc
            except ValueError as exc:
                raise AIServiceError(f"Yandex {operation} вернул не JSON") from exc

        raise AIServiceError(f"Yandex {operation} временно недоступен") from last_error

    @staticmethod
    def _parse_json_object(text: str) -> dict[str, Any]:
        """Извлекает JSON даже если модель случайно добавила ```json."""
        cleaned = text.strip()
        cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
        cleaned = re.sub(r"\s*```$", "", cleaned)

        try:
            value = json.loads(cleaned)
        except json.JSONDecodeError:
            match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
            if not match:
                raise AIServiceError("YandexGPT не вернул JSON с реквизитами")
            try:
                value = json.loads(match.group(0))
            except json.JSONDecodeError as exc:
                raise AIServiceError("Не удалось разобрать JSON от YandexGPT") from exc

        if not isinstance(value, dict):
            raise AIServiceError("YandexGPT вернул JSON неверного типа")
        return value

    def extract_company(self, source_text: str) -> CompanyDetails:
        source_text = source_text.strip()
        if not source_text:
            raise DocumentExtractionError("Нет текста для извлечения реквизитов")

        payload = {
            "modelUri": self.model_uri,
            "completionOptions": {
                "stream": False,
                "temperature": 0.1,
                "maxTokens": 1800,
            },
            "messages": [
                {"role": "system", "text": self.EXTRACTION_PROMPT},
                {
                    "role": "user",
                    "text": "Извлеки реквизиты из текста:\n\n" + source_text,
                },
            ],
        }
        data = self._post_json(
            self.GPT_URL,
            payload,
            operation="YandexGPT",
            timeout=(30, 180),
        )
        try:
            answer = data["result"]["alternatives"][0]["message"]["text"]
        except (KeyError, IndexError, TypeError) as exc:
            raise AIServiceError("Неожиданный формат ответа YandexGPT") from exc

        details = CompanyDetails.from_dict(self._parse_json_object(str(answer)))
        self._restore_number_fields(details, source_text)
        return details

    @staticmethod
    def _restore_number_fields(details: CompanyDetails, source_text: str) -> None:
        """Подстраховывает ИИ: берёт явно подписанные номера из исходного текста."""
        patterns = {
            "inn": r"(?<![\wА-Яа-я])ИНН\s*[:№-]?\s*(\d{10}|\d{12})(?!\d)",
            "checking_account": (
                r"(?:расч[её]тн(?:ый|ого)\s+сч[её]т|р/?с)\s*[:№-]?\s*(\d{20})(?!\d)"
            ),
            "correspondent_account": (
                r"(?:корр(?:еспондентский)?\.?\s*сч[её]т|к/?с)"
                r"\s*[:№-]?\s*(\d{20})(?!\d)"
            ),
            "bik": r"(?<![\wА-Яа-я])БИК\s*[:№-]?\s*(\d{9})(?!\d)",
            "kpp": r"(?<![\wА-Яа-я])КПП\s*[:№-]?\s*(\d{9})(?!\d)",
            "ogrn": r"(?<![\wА-Яа-я])ОГРН(?:ИП)?\s*[:№-]?\s*(\d{13}|\d{15})(?!\d)",
        }
        for field_name, pattern in patterns.items():
            current = str(getattr(details, field_name) or "").strip().casefold()
            if current not in CompanyDetails.EMPTY_VALUES:
                continue
            match = re.search(pattern, source_text, flags=re.IGNORECASE)
            if match:
                setattr(details, field_name, match.group(1))

    @staticmethod
    def _extract_ocr_text(response: dict[str, Any]) -> str:
        """Поддерживает несколько фактических форм ответа Vision OCR."""
        full_texts: list[str] = []
        line_texts: list[str] = []

        def walk(value: Any) -> None:
            if isinstance(value, dict):
                full_text = value.get("fullText")
                if isinstance(full_text, str) and full_text.strip():
                    full_texts.append(full_text.strip())
                text = value.get("text")
                if isinstance(text, str) and text.strip():
                    line_texts.append(text.strip())
                for nested in value.values():
                    walk(nested)
            elif isinstance(value, list):
                for nested in value:
                    walk(nested)

        walk(response)
        if full_texts:
            return max(full_texts, key=len)
        if line_texts:
            # Убираем дубликаты, сохраняя исходный порядок.
            return "\n".join(dict.fromkeys(line_texts))
        raise DocumentExtractionError("Yandex OCR не распознал текст")

    def recognize_jpeg(self, content: bytes) -> str:
        payload = {
            "mimeType": "JPEG",
            "languageCodes": ["ru", "en"],
            "model": "page",
            "content": base64.b64encode(content).decode("ascii"),
        }
        response = self._post_json(
            self.OCR_URL,
            payload,
            operation="OCR",
            timeout=(120, 240),
            attempts=3,
        )
        return self._extract_ocr_text(response)


# ---------------------------------------------------------------------------
# Файлы, OCR и шаблон договора
# ---------------------------------------------------------------------------


class DocumentService:
    """Скачивание вложений MAX и извлечение текста."""

    ALLOWED_DOCUMENT_EXTENSIONS = frozenset({".docx", ".pdf"})
    ALLOWED_IMAGE_EXTENSIONS = frozenset(
        {".jpg", ".jpeg", ".png", ".gif", ".bmp", ".tif", ".tiff", ".webp"}
    )
    MAX_FILE_BYTES = 20 * 1024 * 1024
    MAX_DOCUMENT_CHARS = 80_000
    MAX_SCAN_PAGES = 10
    OCR_PDF_DPI = 200
    OCR_MAX_IMAGE_BYTES = 8 * 1024 * 1024
    OCR_MAX_PIXELS = 12_000_000
    OCR_MIN_PIXELS = 160_000
    OCR_MIN_SIDE = 100

    def __init__(self, config: Config, ai: YandexAI) -> None:
        self.work_dir = config.work_dir
        self.ai = ai
        self.http = requests.Session()
        self.http.trust_env = False

    @staticmethod
    def attachment_url(attachment: dict[str, Any]) -> str | None:
        preferred = ("url", "download_url")

        def walk(value: Any) -> str | None:
            if isinstance(value, dict):
                for key in preferred:
                    candidate = value.get(key)
                    if isinstance(candidate, str) and candidate.startswith(
                        ("http://", "https://")
                    ):
                        return candidate
                for nested in value.values():
                    found = walk(nested)
                    if found:
                        return found
            elif isinstance(value, list):
                for nested in value:
                    found = walk(nested)
                    if found:
                        return found
            return None

        return walk(attachment.get("payload") or attachment)

    @staticmethod
    def attachment_filename(attachment: dict[str, Any]) -> str | None:
        preferred = ("filename", "file_name", "name")

        def walk(value: Any) -> str | None:
            if isinstance(value, dict):
                for key in preferred:
                    candidate = value.get(key)
                    if isinstance(candidate, str) and candidate.strip():
                        return candidate.strip()
                for nested in value.values():
                    found = walk(nested)
                    if found:
                        return found
            elif isinstance(value, list):
                for nested in value:
                    found = walk(nested)
                    if found:
                        return found
            return None

        return walk(attachment.get("payload") or attachment)

    def _detect_suffix(
        self,
        attachment: dict[str, Any],
        url: str,
        response: requests.Response,
        expected_type: str,
    ) -> str:
        accepted = (
            self.ALLOWED_IMAGE_EXTENSIONS
            if expected_type == "image"
            else self.ALLOWED_DOCUMENT_EXTENSIONS
        )

        disposition = response.headers.get("Content-Disposition", "")
        match = re.search(
            r"filename\*?=(?:UTF-8''|\")?([^\";]+)",
            disposition,
            flags=re.IGNORECASE,
        )
        disposition_name = unquote(match.group(1).strip()) if match else ""

        candidates = (
            self.attachment_filename(attachment) or "",
            disposition_name,
            unquote(urlparse(url).path),
        )
        for candidate in candidates:
            suffix = Path(candidate).suffix.lower()
            if suffix in accepted:
                return suffix

        content_type = response.headers.get("Content-Type", "").split(";", 1)[0]
        known = {
            "application/pdf": ".pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
            "image/jpeg": ".jpg",
            "image/png": ".png",
            "image/gif": ".gif",
            "image/bmp": ".bmp",
            "image/tiff": ".tiff",
            "image/webp": ".webp",
        }
        if known.get(content_type) in accepted:
            return known[content_type]

        content = response.content
        if expected_type == "file" and content.lstrip().startswith(b"%PDF-"):
            return ".pdf"
        if expected_type == "file" and content.startswith(b"PK"):
            try:
                with zipfile.ZipFile(BytesIO(content)) as archive:
                    if "word/document.xml" in archive.namelist():
                        return ".docx"
            except zipfile.BadZipFile:
                pass

        signatures = (
            (b"\xff\xd8\xff", ".jpg"),
            (b"\x89PNG\r\n\x1a\n", ".png"),
            (b"GIF87a", ".gif"),
            (b"GIF89a", ".gif"),
            (b"BM", ".bmp"),
            (b"II*\x00", ".tiff"),
            (b"MM\x00*", ".tiff"),
        )
        if expected_type == "image":
            for signature, suffix in signatures:
                if content.startswith(signature):
                    return suffix

        raise DocumentExtractionError(
            "MAX передал вложение без распознаваемого формата"
        )

    def download_attachment(
        self,
        message: dict[str, Any],
        expected_type: str,
    ) -> Path:
        attachments = (message.get("body") or {}).get("attachments") or []
        selected = next(
            (item for item in attachments if item.get("type") == expected_type),
            None,
        )
        if selected is None:
            human = "изображение" if expected_type == "image" else "файл"
            raise DocumentExtractionError(f"Сообщение не содержит {human}")

        url = self.attachment_url(selected)
        if not url:
            raise DocumentExtractionError("MAX не прислал URL вложения")

        response = self.http.get(url, timeout=180)
        response.raise_for_status()
        if len(response.content) > self.MAX_FILE_BYTES:
            raise DocumentExtractionError("Файл превышает лимит 20 МБ")

        suffix = self._detect_suffix(selected, url, response, expected_type)
        target = self.work_dir / f"input_{uuid4().hex}{suffix}"
        target.write_bytes(response.content)
        return target

    @staticmethod
    def safe_delete(path: Path | None) -> None:
        if path is None:
            return
        try:
            path.unlink(missing_ok=True)
        except OSError:
            logging.getLogger("files").exception(
                "Не удалось удалить временный файл %s", path
            )

    def _prepare_ocr_jpeg(self, source: Image.Image) -> bytes:
        image = ImageOps.exif_transpose(source)
        if (
            image.width * image.height < self.OCR_MIN_PIXELS
            or min(image.width, image.height) < self.OCR_MIN_SIDE
        ):
            raise DocumentExtractionError(
                "Изображение слишком маленькое для распознавания"
            )

        if image.mode in {"RGBA", "LA"}:
            background = Image.new("RGB", image.size, "white")
            background.paste(image.convert("RGB"), mask=image.getchannel("A"))
            image = background
        else:
            image = image.convert("RGB")

        pixels = image.width * image.height
        if pixels > self.OCR_MAX_PIXELS:
            scale = (self.OCR_MAX_PIXELS / pixels) ** 0.5
            image = image.resize(
                (
                    max(1, round(image.width * scale)),
                    max(1, round(image.height * scale)),
                ),
                Image.Resampling.LANCZOS,
            )

        result = b""
        for quality in (88, 80, 72, 64, 56):
            buffer = BytesIO()
            image.save(buffer, format="JPEG", quality=quality, optimize=True)
            result = buffer.getvalue()
            if len(result) <= self.OCR_MAX_IMAGE_BYTES:
                return result
        raise DocumentExtractionError("Изображение слишком велико для OCR")

    def extract_image_text(self, path: Path) -> str:
        try:
            with Image.open(path) as image:
                content = self._prepare_ocr_jpeg(image)
        except DocumentExtractionError:
            raise
        except (UnidentifiedImageError, OSError, ValueError) as exc:
            raise DocumentExtractionError(
                "Изображение повреждено или имеет неподдерживаемый формат"
            ) from exc
        return self.ai.recognize_jpeg(content)

    def extract_scan_pdf_text(self, path: Path) -> str:
        pages: list[str] = []
        try:
            document = fitz.open(path)
        except Exception as exc:
            raise DocumentExtractionError("Не удалось открыть PDF") from exc

        with document:
            if document.needs_pass:
                raise DocumentExtractionError("PDF защищён паролем")
            if len(document) == 0:
                raise DocumentExtractionError("PDF не содержит страниц")
            if len(document) > self.MAX_SCAN_PAGES:
                raise DocumentExtractionError(
                    f"В PDF больше {self.MAX_SCAN_PAGES} страниц"
                )

            for index, page in enumerate(document, start=1):
                pixmap = page.get_pixmap(
                    dpi=self.OCR_PDF_DPI,
                    colorspace=fitz.csRGB,
                    alpha=False,
                )
                image = Image.frombytes(
                    "RGB",
                    (pixmap.width, pixmap.height),
                    pixmap.samples,
                )
                text = self.ai.recognize_jpeg(self._prepare_ocr_jpeg(image))
                pages.append(f"--- Страница {index} ---\n{text}")

        return "\n".join(pages)

    def extract_document_text(self, path: Path) -> str:
        suffix = path.suffix.lower()
        if suffix == ".docx":
            try:
                document = Document(path)
            except Exception as exc:
                raise DocumentExtractionError("Не удалось открыть DOCX") from exc

            parts = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append("\t".join(cell.text for cell in row.cells))
            for section in document.sections:
                parts.extend(p.text for p in section.header.paragraphs)
                parts.extend(p.text for p in section.footer.paragraphs)
            text = "\n".join(parts)

        elif suffix == ".pdf":
            try:
                with fitz.open(path) as document:
                    if document.needs_pass:
                        raise DocumentExtractionError("PDF защищён паролем")
                    text = "\n".join(page.get_text("text") for page in document)
            except DocumentExtractionError:
                raise
            except Exception as exc:
                raise DocumentExtractionError("Не удалось открыть PDF") from exc
        else:
            raise DocumentExtractionError(
                "Поддерживаются только DOCX и PDF. Старый DOC сохраните как DOCX."
            )

        normalized = "\n".join(
            line.strip()
            for line in text.replace("\x00", "").splitlines()
            if line.strip()
        )
        if not normalized:
            if suffix == ".pdf":
                raise DocumentExtractionError(
                    "В PDF нет текстового слоя. Выберите кнопку «PDF-скан»."
                )
            raise DocumentExtractionError("Документ не содержит текста")
        if len(normalized) > self.MAX_DOCUMENT_CHARS:
            raise DocumentExtractionError(
                f"Текст документа превышает {self.MAX_DOCUMENT_CHARS} символов"
            )
        return normalized


class CounterService:
    """Безопасно выдаёт последовательные номера договоров и счетов."""

    def __init__(self, directory: Path) -> None:
        self.directory = directory / "counters"
        self.directory.mkdir(parents=True, exist_ok=True)
        self.lock = threading.Lock()

    def next_number(self, prefix: str, kind: str, date_format: str) -> str:
        today = date.today()
        path = self.directory / f"{prefix}_{kind}.txt"

        with self.lock:
            current = 0
            try:
                stored_date, stored_value = path.read_text(
                    encoding="utf-8"
                ).strip().split("|", 1)
                if stored_date == today.isoformat():
                    current = int(stored_value)
            except (OSError, ValueError):
                current = 0

            current += 1
            temporary = path.with_suffix(".tmp")
            temporary.write_text(
                f"{today.isoformat()}|{current}", encoding="utf-8"
            )
            temporary.replace(path)

        return f"{prefix}{datetime.now().strftime(date_format)}-{current}"


class ContractService:
    """Рассчитывает суммы и заполняет Word-шаблон."""

    def __init__(self, config: Config) -> None:
        self.config = config
        self.counters = CounterService(config.work_dir)

    @staticmethod
    def amount_words(value: int) -> str:
        words = num2words(value, lang="ru")
        return words[:1].upper() + words[1:]

    def build(
        self,
        employee_code: str,
        service_cost: int,
        sets_count: int,
        valid_until: str,
        company: CompanyDetails,
    ) -> Path:
        company.validate()

        printing_cost = self.config.print_price_per_set * sets_count
        delivery_cost = self.config.delivery_price if sets_count else 0
        documents_total = printing_cost + delivery_cost
        contract_total = service_cost + documents_total

        contract_number = self.counters.next_number(
            employee_code, "contracts", "%d%m%y"
        )
        invoice_number = self.counters.next_number(
            employee_code, "invoices", "%d%m"
        )

        if company.is_individual_entrepreneur:
            basis = "Листа записи ЕГРИП"
        else:
            basis = "Устава"

        customer_intro = (
            f"{company.display_name}, именуемое в дальнейшем «Заказчик», "
            f"в лице {company.director_title_genitive} "
            f"{company.director_name_genitive}, действующего на основании {basis}"
        )

        bank_details = "\n".join(
            [
                company.bank_name,
                f"р/с {company.checking_account}",
                f"к/с {company.correspondent_account}",
                f"БИК {company.bik}",
            ]
        )
        registration_details = "\n".join(
            item
            for item in [
                company.legal_address,
                f"ИНН {company.inn}",
                f"КПП {company.kpp}" if company.kpp else "",
                f"ОГРН/ОГРНИП {company.ogrn}",
            ]
            if item
        )

        now = datetime.now()
        context = {
            "contract_number": contract_number,
            "invoice_number": invoice_number,
            "date_start": f"{now.day:02d}.{now.month:02d}.{now.year}",
            "date_end": valid_until,
            "customer_name": company.display_name,
            "customer_intro": customer_intro,
            "service_cost": service_cost,
            "service_cost_words": self.amount_words(service_cost),
            "sets_count": sets_count,
            "printing_cost": printing_cost,
            "delivery_cost": delivery_cost,
            "documents_total": documents_total,
            "documents_total_words": self.amount_words(documents_total),
            "contract_total": contract_total,
            "contract_total_words": self.amount_words(contract_total),
            "director_title": company.director_title_nominative,
            "director_name": company.director_name_nominative,
            "inn": company.inn,
            "kpp": company.kpp,
            "ogrn": company.ogrn,
            "legal_address": company.legal_address,
            "bank_details": bank_details,
            "registration_details": registration_details,
            "company_json": json.dumps(asdict(company), ensure_ascii=False, indent=2),
        }

        document = DocxTemplate(self.config.template_path)
        document.render(context)
        output = self.config.work_dir / f"ДОГОВОР_{contract_number}.docx"
        document.save(output)
        return output


# ---------------------------------------------------------------------------
# Сценарий общения с пользователем
# ---------------------------------------------------------------------------

START_BUTTON_TEXT = "🚀 Начать"
RESTART_BUTTON_TEXT = "🔄 Перезапустить"
START_WORDS = {
    "/start",
    START_BUTTON_TEXT.casefold(),
    RESTART_BUTTON_TEXT.casefold(),
    "старт",
    "начать",
    "начало",
    "начни",
    "перезапустить",
    "начать заново",
}


@dataclass(slots=True)
class UserSession:
    stage: str = "idle"
    employee_code: str = ""
    service_cost: int = 0
    sets_count: int = 0
    valid_until: str = ""
    scenario_id: str = field(default_factory=lambda: uuid4().hex)


class ContractBotApp:
    """Маршрутизация обновлений и конечный автомат диалога."""

    def __init__(self, config: Config) -> None:
        self.config = config
        self.max = MaxClient(config)
        self.ai = YandexAI(config)
        self.documents = DocumentService(config, self.ai)
        self.contracts = ContractService(config)
        self.sessions: dict[int, UserSession] = {}
        self.recent_actions: dict[int, tuple[str, float]] = {}
        self.logger = logging.getLogger("max-contract-bot")

    @staticmethod
    def control_keyboard() -> dict[str, Any]:
        return {
            "type": "inline_keyboard",
            "payload": {
                "buttons": [
                    [
                        {
                            "type": "callback",
                            "text": START_BUTTON_TEXT,
                            "payload": "start",
                        },
                        {
                            "type": "callback",
                            "text": RESTART_BUTTON_TEXT,
                            "payload": "restart",
                        },
                    ]
                ]
            },
        }

    @staticmethod
    def source_keyboard() -> dict[str, Any]:
        return {
            "type": "inline_keyboard",
            "payload": {
                "buttons": [
                    [
                        {"type": "callback", "text": "🧾 DOCX/PDF", "payload": "doc"},
                        {"type": "callback", "text": "🖼️ Картинка", "payload": "pic"},
                    ],
                    [
                        {"type": "callback", "text": "📄 PDF-скан", "payload": "pdf"},
                        {"type": "callback", "text": "📢 Текст", "payload": "mes"},
                    ],
                    [
                        {
                            "type": "callback",
                            "text": RESTART_BUTTON_TEXT,
                            "payload": "restart",
                        }
                    ],
                ]
            },
        }

    def allowed(self, user_id: int) -> bool:
        if user_id in self.config.staff_ids:
            return True
        self.logger.warning("Запрещённый доступ | user_id=%s", user_id)
        try:
            self.max.send(
                user_id,
                "❌ <b>Вам запрещён доступ к боту.</b> Обратитесь к администратору.",
            )
        except Exception:
            self.logger.exception("Не удалось отправить отказ | user_id=%s", user_id)
        return False

    def is_repeated(self, user_id: int, action: str) -> bool:
        now = time.monotonic()
        previous, previous_at = self.recent_actions.get(user_id, ("", 0.0))
        self.recent_actions[user_id] = (action, now)
        return previous == action and now - previous_at < 1.5

    @staticmethod
    def message_text(message: dict[str, Any]) -> str:
        return str((message.get("body") or {}).get("text") or "").strip()

    @staticmethod
    def sender_id(message: dict[str, Any], update: dict[str, Any]) -> int:
        sender = message.get("sender") or update.get("user") or {}
        value = sender.get("user_id")
        if value is None:
            raise ValueError("MAX update не содержит user_id отправителя")
        return int(value)

    def hello_text(self) -> str:
        if self.config.hello_file.is_file():
            return self.config.hello_file.read_text(
                encoding="utf-8", errors="replace"
            ).strip()
        return "Здравствуйте! Введите фамилию сотрудника, который формирует договор."

    def show_start_menu(self, user_id: int, text: str | None = None) -> None:
        if not self.allowed(user_id):
            return
        self.sessions[user_id] = UserSession()
        self.max.send(
            user_id,
            text or "Нажмите кнопку, чтобы сформировать новый проект договора.",
            [self.control_keyboard()],
        )

    def start_scenario(self, user_id: int) -> None:
        if not self.allowed(user_id):
            return
        self.sessions[user_id] = UserSession(stage="name")
        self.max.send(user_id, self.hello_text(), [self.control_keyboard()])

    def _session(self, user_id: int) -> UserSession:
        return self.sessions.setdefault(user_id, UserSession())

    def _send_missing_fields(
        self,
        user_id: int,
        error: MissingCompanyDetailsError,
    ) -> None:
        fields = "\n".join(f"• <b>{item}</b>" for item in error.fields)
        self.max.send(
            user_id,
            "⚠️ Договор пока нельзя сформировать. Не хватает реквизитов:\n\n"
            f"{fields}\n\nОтправьте исправленные реквизиты повторно.",
            [self.control_keyboard()],
        )

    @staticmethod
    def _attachment_pending(error: Exception) -> bool:
        """Проверяет известные варианты ошибки ещё не обработанного файла."""
        text = str(error).casefold()
        return any(
            marker in text
            for marker in (
                "attachment.not.ready",
                "attachment.file.not.processed",
                "errors.process.attachment.file.not.processed",
                "not processed",
            )
        )

    def _send_contract_file(self, user_id: int, path: Path) -> None:
        attachment = self.max.upload_file(path)
        delays = (1, 2, 4, 8, 12, 20)

        for attempt, delay in enumerate(delays, start=1):
            try:
                self.max.send(user_id, "✅ Проект договора готов.", [attachment])
                return
            except (MaxAPIError, requests.HTTPError) as exc:
                if not self._attachment_pending(exc):
                    raise
                if attempt == len(delays):
                    break
                self.logger.info(
                    "Файл ещё обрабатывается MAX | user_id=%s | attempt=%s/%s",
                    user_id,
                    attempt,
                    len(delays),
                )
                time.sleep(delay)

        raise MaxAPIError("MAX не успел подготовить загруженный договор")

    def _finish_document(self, user_id: int, company: CompanyDetails) -> None:
        current = self._session(user_id)
        company.validate()

        self.max.send(
            user_id,
            "🧾 Реквизиты извлечены. Рассчитываю суммы и заполняю шаблон…",
            [self.control_keyboard()],
        )
        output: Path | None = None
        try:
            output = self.contracts.build(
                employee_code=current.employee_code,
                service_cost=current.service_cost,
                sets_count=current.sets_count,
                valid_until=current.valid_until,
                company=company,
            )
            self.max.send(
                user_id,
                "📤 Договор сформирован. Загружаю файл в MAX…",
                [self.control_keyboard()],
            )
            self._send_contract_file(user_id, output)
        finally:
            self.documents.safe_delete(output)

        self.show_start_menu(
            user_id,
            "✅ Договор отправлен. Можно сформировать следующий.",
        )

    def handle_text_or_attachment(
        self,
        user_id: int,
        message: dict[str, Any],
    ) -> None:
        if not self.allowed(user_id):
            return

        current = self._session(user_id)
        text = self.message_text(message)
        normalized = text.casefold()

        if normalized in START_WORDS:
            self.start_scenario(user_id)
            return

        if current.stage == "idle":
            self.show_start_menu(
                user_id,
                "⚠️ Активный сценарий не найден. Нажмите «Начать».",
            )
            return

        if current.stage == "name":
            letters = [char for char in text if char.isalpha()]
            if len(letters) < 2:
                self.max.send(user_id, "❌ Введите фамилию минимум из двух букв.")
                return
            current.employee_code = "".join(letters[:2]).upper()
            ending = datetime.now() + relativedelta(years=1)
            month_names = (
                "",
                "января",
                "февраля",
                "марта",
                "апреля",
                "мая",
                "июня",
                "июля",
                "августа",
                "сентября",
                "октября",
                "ноября",
                "декабря",
            )
            current.valid_until = (
                f"{ending.day} {month_names[ending.month]} {ending.year} года"
            )
            current.stage = "cost"
            self.max.send(
                user_id,
                f"✅ Код документов: <b>{current.employee_code}</b>\n"
                "❔ Введите стоимость разовой услуги целым числом.",
                [self.control_keyboard()],
            )
            return

        if current.stage == "cost":
            if not text.isdigit():
                self.max.send(user_id, "❌ Введите стоимость целым числом без пробелов.")
                return
            current.service_cost = int(text)
            current.stage = "sets"
            self.max.send(
                user_id,
                f"✅ Стоимость: <b>{current.service_cost:,} ₽</b>\n"
                "❔ Введите количество печатных комплектов.".replace(",", " "),
                [self.control_keyboard()],
            )
            return

        if current.stage == "sets":
            if not text.isdigit():
                self.max.send(user_id, "❌ Введите количество целым числом.")
                return
            current.sets_count = int(text)
            current.stage = "choice"
            self.max.send(
                user_id,
                f"✅ Комплектов: <b>{current.sets_count}</b>\n"
                "❔ Выберите источник реквизитов.",
                [self.source_keyboard()],
            )
            return

        source: Path | None = None
        try:
            if current.stage == "doc":
                source = self.documents.download_attachment(message, "file")
                self.max.send(
                    user_id,
                    "✅ Файл получен. Извлекаю текст и реквизиты…",
                    [self.control_keyboard()],
                )
                text_from_file = self.documents.extract_document_text(source)
                self._finish_document(
                    user_id, self.ai.extract_company(text_from_file)
                )

            elif current.stage == "pdf":
                source = self.documents.download_attachment(message, "file")
                if source.suffix.lower() != ".pdf":
                    raise DocumentExtractionError("Для этого шага отправьте PDF")
                self.max.send(
                    user_id,
                    "✅ PDF-скан получен. Распознаю страницы…",
                    [self.control_keyboard()],
                )
                scanned_text = self.documents.extract_scan_pdf_text(source)
                self._finish_document(
                    user_id, self.ai.extract_company(scanned_text)
                )

            elif current.stage == "pic":
                source = self.documents.download_attachment(message, "image")
                self.max.send(
                    user_id,
                    "✅ Изображение получено. Распознаю текст…",
                    [self.control_keyboard()],
                )
                image_text = self.documents.extract_image_text(source)
                self._finish_document(user_id, self.ai.extract_company(image_text))

            elif current.stage == "mes":
                if not text:
                    raise DocumentExtractionError("Отправьте текст с реквизитами")
                self.max.send(
                    user_id,
                    "✅ Сообщение получено. Извлекаю реквизиты…",
                    [self.control_keyboard()],
                )
                self._finish_document(user_id, self.ai.extract_company(text))

            elif current.stage == "choice":
                self.max.send(user_id, "❔ Сначала выберите источник кнопкой.")
            else:
                self.show_start_menu(user_id)

        except MissingCompanyDetailsError as exc:
            self.logger.warning(
                "Не хватает реквизитов | user_id=%s | fields=%s",
                user_id,
                ", ".join(exc.fields),
            )
            self._send_missing_fields(user_id, exc)
        except (DocumentExtractionError, AIServiceError, MaxAPIError) as exc:
            self.logger.warning(
                "Обработка отклонена | user_id=%s | error=%s", user_id, exc
            )
            self.max.send(
                user_id,
                f"❌ {exc}\n\nИсправьте данные и повторите текущий шаг.",
                [self.control_keyboard()],
            )
        except requests.RequestException as exc:
            self.logger.exception("Сетевая ошибка | user_id=%s", user_id)
            self.max.send(
                user_id,
                "❌ Не удалось скачать или отправить файл. Проверьте интернет.",
                [self.control_keyboard()],
            )
        except Exception:
            self.logger.exception("Неожиданная ошибка | user_id=%s", user_id)
            self.max.send(
                user_id,
                "❌ Произошла внутренняя ошибка. Подробности записаны в лог.",
                [self.control_keyboard()],
            )
        finally:
            self.documents.safe_delete(source)

    def handle_callback(self, update: dict[str, Any]) -> None:
        callback = update.get("callback") or {}
        message = update.get("message") or callback.get("message") or {}
        user = (
            callback.get("user")
            or update.get("user")
            or message.get("sender")
            or {}
        )
        user_id = int(user["user_id"])
        if not self.allowed(user_id):
            return

        payload = str(callback.get("payload") or "")
        if self.is_repeated(user_id, f"callback:{payload}"):
            return
        if payload in {"start", "restart"}:
            self.start_scenario(user_id)
            return

        current = self._session(user_id)
        if current.stage != "choice":
            self.show_start_menu(
                user_id,
                "⚠️ Предыдущий сценарий завершён. Нажмите «Начать».",
            )
            return

        choices = {
            "doc": ("doc", "🧾 Отправьте DOCX или PDF с текстовым слоем."),
            "pdf": ("pdf", "📄 Отправьте PDF-скан."),
            "pic": ("pic", "🖼️ Отправьте изображение реквизитов."),
            "mes": ("mes", "📢 Отправьте реквизиты текстовым сообщением."),
        }
        selected = choices.get(payload)
        if selected:
            current.stage = selected[0]
            self.max.send(user_id, selected[1], [self.control_keyboard()])

    def handle_update(self, update: dict[str, Any]) -> None:
        update_type = update.get("update_type")
        if update_type == "bot_started":
            user = update.get("user") or {}
            self.show_start_menu(int(user["user_id"]))
            return
        if update_type == "message_callback":
            self.handle_callback(update)
            return
        if update_type == "message_created":
            message = update.get("message") or {}
            user_id = self.sender_id(message, update)
            text = self.message_text(message).casefold()
            if text in START_WORDS and self.is_repeated(user_id, f"message:{text}"):
                return
            self.handle_text_or_attachment(user_id, message)

    def run(self) -> None:
        identity = self.max.me()
        self.logger.info(
            "MAX authentication OK | bot_id=%s | username=%s",
            identity.get("user_id"),
            identity.get("username"),
        )

        for user_id in self.config.staff_ids:
            try:
                self.show_start_menu(
                    user_id,
                    "✅ Бот запущен. Можно сформировать новый договор.",
                )
            except Exception:
                self.logger.warning(
                    "Не удалось уведомить сотрудника | user_id=%s", user_id
                )

        marker: int | None = None
        self.logger.info("MAX contract bot started in long polling mode")

        while True:
            try:
                page = self.max.updates(marker)
                next_marker = page.get("marker", marker)
                for update in page.get("updates", []):
                    try:
                        self.handle_update(update)
                    except Exception:
                        self.logger.exception(
                            "Update skipped | update_type=%s",
                            update.get("update_type"),
                        )
                marker = next_marker
            except MaxAuthenticationError:
                self.logger.critical("MAX отклонил токен. Бот остановлен.")
                return
            except (MaxAPIError, requests.RequestException):
                self.logger.exception("Ошибка long polling; повтор через 3 секунды")
                time.sleep(3)
            except KeyboardInterrupt:
                self.logger.info("Бот остановлен пользователем")
                return
            except Exception:
                self.logger.exception("Необработанная ошибка цикла; повтор через 3 секунды")
                time.sleep(3)


def main() -> None:
    config = Config.load()
    configure_logging(config)
    ContractBotApp(config).run()


if __name__ == "__main__":
    main()
